"""文档输出模块"""

import os
from datetime import datetime
from typing import Dict, Optional
from pathlib import Path


class DocumentExporter:
    """文档导出器"""
    
    def __init__(self, config: Dict):
        self.config = config
        self.output_dir = config.get("output_dir", "./output")
        self.formats = config.get("formats", ["markdown", "docx"])
        
        # 确保输出目录存在
        os.makedirs(self.output_dir, exist_ok=True)
    
    def export(self, content: str, filename_template: str, mode: str = "full") -> Dict[str, str]:
        """
        导出文档
        
        Args:
            content: 文档内容
            filename_template: 文件名模板（支持 {timestamp} 占位符）
            mode: 导出模式（full 或 summary）
            
        Returns:
            导出文件路径字典 {format: filepath}
        """
        # 生成文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = filename_template.format(timestamp=timestamp)
        base_name = os.path.splitext(base_filename)[0]
        
        exported_files = {}
        
        # 导出各种格式
        if "markdown" in self.formats:
            md_path = self._export_markdown(content, base_name)
            if md_path:
                exported_files["markdown"] = md_path
        
        if "docx" in self.formats:
            docx_path = self._export_docx(content, base_name)
            if docx_path:
                exported_files["docx"] = docx_path
        
        if "pdf" in self.formats:
            pdf_path = self._export_pdf(content, base_name)
            if pdf_path:
                exported_files["pdf"] = pdf_path
        
        return exported_files
    
    def _export_markdown(self, content: str, base_name: str) -> Optional[str]:
        """导出 Markdown"""
        try:
            filename = f"{base_name}.md"
            filepath = os.path.join(self.output_dir, filename)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return filepath
        except Exception as e:
            print(f"导出 Markdown 失败: {str(e)}")
            return None
    
    def _export_docx(self, content: str, base_name: str) -> Optional[str]:
        """导出 Word 文档"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            import re
            
            doc = Document()
            
            # 设置默认字体
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            font.size = Pt(12)
            
            # 解析内容
            lines = content.split('\n')
            current_paragraph = None
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    # 空行
                    if current_paragraph:
                        current_paragraph = None
                    continue
                
                # 检查是否是标题（讲话人标记或一级标题）
                if line.startswith('【') and line.endswith('】'):
                    # 讲话人标记，作为标题
                    p = doc.add_heading(line, level=2)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    current_paragraph = None
                
                elif re.match(r'^[一二三四五六七八九十]+[、.]', line) or re.match(r'^\d+[、.]', line):
                    # 编号标题
                    p = doc.add_heading(line, level=2)
                    current_paragraph = None
                
                elif line.startswith('# '):
                    # Markdown 一级标题
                    p = doc.add_heading(line[2:], level=1)
                    current_paragraph = None
                
                elif line.startswith('## '):
                    # Markdown 二级标题
                    p = doc.add_heading(line[3:], level=2)
                    current_paragraph = None
                
                elif line.startswith('### '):
                    # Markdown 三级标题
                    p = doc.add_heading(line[4:], level=3)
                    current_paragraph = None
                
                else:
                    # 普通段落
                    if current_paragraph is None:
                        current_paragraph = doc.add_paragraph()
                    else:
                        # 同一段落内换行
                        current_paragraph.add_run('\n')
                    
                    current_paragraph.add_run(line)
            
            # 保存
            filename = f"{base_name}.docx"
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)
            
            return filepath
        
        except ImportError:
            print("警告: 未安装 python-docx，跳过 Word 导出")
            return None
        except Exception as e:
            print(f"导出 Word 失败: {str(e)}")
            return None
    
    def _export_pdf(self, content: str, base_name: str) -> Optional[str]:
        """导出 PDF 文档"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
            import re
            
            filename = f"{base_name}.pdf"
            filepath = os.path.join(self.output_dir, filename)
            
            # 创建 PDF 文档
            doc = SimpleDocTemplate(
                filepath,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # 样式
            styles = getSampleStyleSheet()
            
            # 尝试注册中文字体（如果可用）
            try:
                # 尝试使用系统字体
                from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
                chinese_font = 'STSong-Light'
            except:
                chinese_font = 'Helvetica'
            
            # 自定义样式
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName=chinese_font,
                fontSize=12,
                leading=18,
                alignment=TA_LEFT,
                spaceAfter=6
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontName=chinese_font,
                fontSize=14,
                leading=20,
                alignment=TA_LEFT,
                spaceAfter=12,
                spaceBefore=12
            )
            
            # 构建内容
            story = []
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    story.append(Spacer(1, 6))
                    continue
                
                # 检查是否是标题
                if line.startswith('【') and line.endswith('】'):
                    # 讲话人标记
                    story.append(Paragraph(line, heading_style))
                    story.append(Spacer(1, 6))
                
                elif re.match(r'^[一二三四五六七八九十]+[、.]', line) or re.match(r'^\d+[、.]', line):
                    # 编号标题
                    story.append(Paragraph(line, heading_style))
                    story.append(Spacer(1, 6))
                
                elif line.startswith('# '):
                    # Markdown 一级标题
                    story.append(Paragraph(line[2:], heading_style))
                    story.append(Spacer(1, 6))
                
                elif line.startswith('## '):
                    # Markdown 二级标题
                    story.append(Paragraph(line[3:], heading_style))
                    story.append(Spacer(1, 6))
                
                elif line.startswith('### '):
                    # Markdown 三级标题
                    story.append(Paragraph(line[4:], heading_style))
                    story.append(Spacer(1, 6))
                
                else:
                    # 普通段落
                    # 转义 HTML 特殊字符
                    line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(line, normal_style))
            
            # 生成 PDF
            doc.build(story)
            
            return filepath
        
        except ImportError:
            print("警告: 未安装 reportlab，跳过 PDF 导出")
            return None
        except Exception as e:
            print(f"导出 PDF 失败: {str(e)}")
            return None

